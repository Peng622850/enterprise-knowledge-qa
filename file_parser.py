# file_parser.py
import io

import docx
import pypdf

from logger_setup import get_logger
logger = get_logger(__name__)

SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt"}


def parse_file(filename: str, file_bytes: bytes) -> tuple[str, list[dict]]:
    """
    返回 (全文文本, chunk_metadatas)
    chunk_metadatas 里每个元素对应一个逻辑块的来源信息
    """
    ext = filename.lower().split(".")[-1]
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持的文件格式：{ext}，请上传 pdf / docx / txt")

    if ext == "pdf":
        return parse_pdf(filename, file_bytes)
    elif ext == "docx":
        return parse_docx(filename, file_bytes)
    else:
        text = file_bytes.decode("utf-8", errors="ignore")
        return text, [{"source": filename, "page": 1, "block_type": "text"}]


def parse_pdf(filename: str, file_bytes: bytes) -> tuple[str, list[dict]]:
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    blocks = []
    metadatas = []

    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text()
        if text and text.strip():
            blocks.append(text.strip())
            metadatas.append({
                "source": filename,
                "page": page_num,
                "block_type": "text",
            })

    logger.info(f"PDF 解析完成：{filename}，共 {len(reader.pages)} 页，{len(blocks)} 个文本块")
    return "\n\n".join(blocks), metadatas


def parse_docx(filename: str, file_bytes: bytes) -> tuple[str, list[dict]]:
    doc = docx.Document(io.BytesIO(file_bytes))
    blocks = []
    metadatas = []

    # 提取段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        blocks.append(text)
        metadatas.append({
            "source": filename,
            "page": 1,
            "block_type": "paragraph",
            "style": para.style.name,
        })

    # 提取表格，转成 markdown 格式
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if not rows:
            continue

        # 加表头分隔线
        header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
        rows.insert(1, header_sep)

        table_text = "\n".join(rows)
        blocks.append(table_text)
        metadatas.append({
            "source": filename,
            "page": 1,
            "block_type": "table",
            "table_index": t_idx,
        })
        logger.info(f"提取表格 {t_idx+1}：{len(table.rows)} 行")

    logger.info(f"Word 解析完成：{filename}，{len(blocks)} 个块（含表格）")
    return "\n\n".join(blocks), metadatas